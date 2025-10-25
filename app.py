import streamlit as st
import pandas as pd
import zipfile
import os
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Function to extract text from files
def read_text_from_file(file_path):
    ext = file_path.split('.')[-1].lower()
    try:
        if ext == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif ext == 'docx':
            from docx import Document
            doc = Document(file_path)
            return '\n'.join([p.text for p in doc.paragraphs])
        elif ext == 'pdf':
            import fitz
            text = ''
            with fitz.open(file_path) as pdf:
                for page in pdf:
                    text += page.get_text()
            return text
    except:
        return ''
    return ''

# Function to calculate similarity
def rank_resumes(job_description_text, resume_texts, resume_names):
    # If some resumes are empty, they should still be included with score 0
    documents = [job_description_text] + resume_texts
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(documents)
    
    similarity_scores = []
    for i, text in enumerate(resume_texts):
        if text.strip() == '':
            similarity_scores.append(0.0)  # Assign 0 if text is empty
        else:
            sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[i+1:i+2])[0][0]
            similarity_scores.append(sim)
    
    df = pd.DataFrame({
        'Resume': resume_names,
        'Similarity Score': similarity_scores
    }).sort_values(by='Similarity Score', ascending=False)
    return df

# === STREAMLIT UI ===
st.title("📄 AI Resume Screener")

jd_file = st.file_uploader("Upload Job Description", type=['txt', 'docx', 'pdf'])
zip_file = st.file_uploader("Upload Resume Folder (ZIP File)", type=['zip'])

if jd_file and zip_file:
    # Read Job Description
    jd_ext = jd_file.name.split('.')[-1].lower()
    if jd_ext == 'txt':
        job_description_text = jd_file.read().decode('utf-8')
    elif jd_ext == 'docx':
        from docx import Document
        doc = Document(jd_file)
        job_description_text = '\n'.join([p.text for p in doc.paragraphs])
    elif jd_ext == 'pdf':
        import fitz
        job_description_text = ""
        with fitz.open(stream=jd_file.read(), filetype='pdf') as pdf:
            for page in pdf:
                job_description_text += page.get_text()

    # Extract ZIP
    extract_path = "uploaded_resumes"
    if not os.path.exists(extract_path):
        os.makedirs(extract_path)

    with zipfile.ZipFile(zip_file, 'r') as zip_ref:
        zip_ref.extractall(extract_path)

    # Read resumes
    resume_texts = []
    resume_names = []
    for root, _, files in os.walk(extract_path):
        for file in files:
            file_path = os.path.join(root, file)
            text = read_text_from_file(file_path)
            # Append all files even if empty
            resume_texts.append(text)
            resume_names.append(file)

    # Rank resumes
    if resume_texts:
        results_df = rank_resumes(job_description_text, resume_texts, resume_names)
        st.success("✅ Resumes ranked successfully!")
        st.dataframe(results_df.style.format({"Similarity Score": "{:.2%}"}), use_container_width=True)

        # Enable CSV download
        csv = results_df.to_csv(index=False).encode('utf-8')
        st.download_button("📥 Download results as CSV", csv, "ranked_resumes.csv", "text/csv")
    else:
        st.warning("No valid resumes found in ZIP file.")
else:
    st.info("Please upload a Job Description and a ZIP file containing resumes.")
